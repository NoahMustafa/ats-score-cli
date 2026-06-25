"""Dump a resume (PDF/DOCX) to clean, structured Markdown for an LLM to read.

Reuses the same extraction as scoring (two-column de-scramble, drawn-bullet
reconstruction, link recovery, glyph cleanup), then maps it to Markdown:
first content line -> `# title`, section headers -> `## heading`, bullets ->
`- item`. Hyperlinks the parser recovered (often hidden behind contact icons)
are appended as a `## Links` section so the model sees them too.
"""

from __future__ import annotations

from pathlib import Path

from .extract import extract, Document
from .checks_ats import _is_heading
from .checks_content import _BULLET


def to_markdown(doc: Document) -> str:
    out: list[str] = []
    title_done = False
    for ln in doc.text.splitlines():
        s = ln.strip()
        if not s:
            continue
        if _BULLET.match(ln):
            out.append("- " + _BULLET.sub("", ln).strip())
        elif not title_done:
            out.append("# " + s)        # first real line is the name/title
            title_done = True
        elif _is_heading(s):
            out += ["", "## " + s, ""]
        else:
            out.append(s)
    md = "\n".join(out).strip() + "\n"
    if doc.links:
        md += "\n## Links\n\n" + "\n".join(f"- {u}" for u in doc.links) + "\n"
    return md


def _resolve_out(resume: Path, out_path: str | Path | None, force: bool) -> Path:
    """Pick the output path. Default: resume name with .md, beside the resume.

    A directory (or trailing slash) -> <resume-stem>.md inside it. A path with
    no suffix gets .md. Without --force, an existing target is auto-renamed
    `name-1.md`, `name-2.md`, … so we never silently clobber.
    """
    if out_path is None:
        target = resume.with_suffix(".md")
    else:
        out = Path(out_path)
        if out.is_dir() or str(out_path).endswith(("/", "\\")):
            target = out / (resume.stem + ".md")
        elif out.suffix:
            target = out
        else:
            target = out.with_suffix(".md")

    if force or not target.exists():
        return target
    n = 1
    while True:
        cand = target.with_name(f"{target.stem}-{n}{target.suffix}")
        if not cand.exists():
            return cand
        n += 1


def export_markdown(resume_path: str | Path, out_path: str | Path | None = None,
                    force: bool = False) -> Path:
    """Extract the resume and write structured Markdown. Returns the path written."""
    resume = Path(resume_path)
    doc = extract(resume)                       # validates suffix / existence
    target = _resolve_out(resume, out_path, force)
    target.parent.mkdir(parents=True, exist_ok=True)
    target.write_text(to_markdown(doc), encoding="utf-8")
    return target


def _selfcheck() -> None:
    import tempfile
    import docx

    with tempfile.TemporaryDirectory() as d:
        p = Path(d) / "r.docx"
        doc = docx.Document()
        doc.add_paragraph("Jane Doe")
        doc.add_paragraph("EXPERIENCE")
        doc.add_paragraph("• Built data pipelines")
        doc.save(str(p))

        md = to_markdown(extract(p))
        assert md.startswith("# Jane Doe"), md
        assert "## EXPERIENCE" in md, md
        assert "- Built data pipelines" in md, md

        # First write lands on r.md; second auto-renames to r-1.md (no clobber).
        a = export_markdown(p)
        assert a == Path(d) / "r.md", a
        b = export_markdown(p)
        assert b == Path(d) / "r-1.md", b
        # --force overwrites in place.
        c = export_markdown(p, force=True)
        assert c == Path(d) / "r.md", c

        # Explicit dir target -> stem.md inside it.
        sub = Path(d) / "out"
        sub.mkdir()
        e = export_markdown(p, sub)
        assert e == sub / "r.md", e

    print("tomd selfcheck ok")


if __name__ == "__main__":
    _selfcheck()
