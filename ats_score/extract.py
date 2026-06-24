"""Resume text + layout extraction for PDF and DOCX.

Returns plain text plus the layout flags the ATS-readiness check needs
(tables, columns, images, scanned). Dispatch is by file suffix.
"""

from __future__ import annotations

import re
from dataclasses import dataclass, field
from pathlib import Path

# Unmappable glyphs pdfplumber emits for icon/symbol fonts.
_CID = re.compile(r"\(cid:\d+\)")
# Zero-width / BOM chars (templates like Enhancv inject these; they silently
# break keyword and spell matching, e.g. a zero-width space inside "JavaScript").
_ZEROWIDTH = re.compile("[​‌‍⁠﻿]")
# Hyphen + newline between letters = a word wrapped across a line break.
_HYPHENBREAK = re.compile(r"(?<=[a-zA-Z])-\n(?=[a-zA-Z])")


@dataclass
class Document:
    text: str
    fmt: str                       # "pdf" or "docx"
    source: Path
    has_tables: bool = False
    has_columns: bool = False
    has_images: bool = False
    is_scanned: bool = False       # PDF with images but ~no extractable text
    drawn_bullets: bool = False    # bullets are vector graphics, not text (ATS risk)
    links: list[str] = field(default_factory=list)   # hyperlink URIs (mailto/tel/http)
    warnings: list[str] = field(default_factory=list)


def _normalize(text: str) -> str:
    # Unify line endings so every downstream check sees \n only, and drop
    # unmappable icon-font glyphs that would otherwise read as spell errors.
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    text = _CID.sub("", text)
    text = _ZEROWIDTH.sub("", text)
    # Pull a hyphen-line-break back onto one line, keeping the hyphen:
    # "third-\nparty" -> "third-party" (real compound preserved), "cor-\nporate"
    # -> "cor-porate" (skipped by the speller). Dropping the hyphen instead would
    # wrongly fuse genuine compounds ("third-party" -> "thirdparty").
    text = _HYPHENBREAK.sub("-", text)
    return text.strip()


def _extract_pdf(path: Path) -> Document:
    import pdfplumber

    parts: list[str] = []
    links: list[str] = []
    has_tables = has_images = has_columns = drawn_bullets = False
    with pdfplumber.open(path) as pdf:
        for page in pdf.pages:
            text, multicol, drawn = _page_text(page)
            parts.append(text)
            if page.extract_tables():
                has_tables = True
            if page.images:
                has_images = True
            if multicol:
                has_columns = True
            if drawn:
                drawn_bullets = True
            links.extend(_pdf_links(page))

    text = _normalize("\n".join(parts))
    # ponytail: scanned heuristic = images present but almost no text. Naive but
    # catches the common "exported scan" case; OCR is out of scope.
    is_scanned = has_images and len(text) < 100
    return Document(
        text=text, fmt="pdf", source=path,
        has_tables=has_tables, has_columns=has_columns,
        has_images=has_images, is_scanned=is_scanned,
        drawn_bullets=drawn_bullets, links=_dedup(links),
    )


def _dedup(items: list[str]) -> list[str]:
    seen: dict[str, None] = {}
    for i in items:
        if i:
            seen.setdefault(i, None)
    return list(seen)


def _pdf_links(page) -> list[str]:
    out: list[str] = []
    for h in (page.hyperlinks or []):
        if h.get("uri"):
            out.append(h["uri"])
    for a in (page.annots or []):
        data = a.get("data")
        action = data.get("A") if isinstance(data, dict) else None
        uri = action.get("URI") if isinstance(action, dict) else None
        if uri:
            out.append(uri.decode() if isinstance(uri, bytes) else uri)
    return out


def _page_text(page) -> tuple[str, bool, bool]:
    """Extract page text, de-scrambling a two-column layout if one is detected.

    Returns (text, is_multicolumn, has_drawn_bullets). For a single column we
    read normally; for two columns we read the left column fully, then the
    right, so the output isn't the line-by-line merge an ATS would garble.
    """
    # ponytail: x_tolerance=2 recovers lost word spacing without over-splitting.
    gutter = _gutter(page)
    if gutter is None:
        text = page.extract_text(x_tolerance=2) or ""
        # Word-exported resumes draw list bullets as vector dots, not text, so
        # extract_text yields no "•". Reconstruct them so bullets are gradable,
        # and flag it: a real ATS won't see drawn bullets as list items.
        dots = _bullet_dots(page)
        drawn = bool(dots) and "•" not in text
        if drawn:
            text = _text_with_bullets(page, dots)
        return text, False, drawn
    left = page.crop((0, 0, gutter, page.height)).extract_text(x_tolerance=2) or ""
    right = page.crop((gutter, 0, page.width, page.height)).extract_text(x_tolerance=2) or ""
    return (left + "\n" + right), True, False


def _bullet_dots(page) -> list[tuple[float, float]]:
    """(x, top) of small left-side vector marks used as list bullets."""
    dots: list[tuple[float, float]] = []
    for o in page.curves + page.rects:
        w = o["x1"] - o["x0"]
        h = o["bottom"] - o["top"]
        if w < 12 and h < 12 and o["x0"] < page.width * 0.5:
            dots.append((o["x0"], o["top"]))
    return dots


def _text_with_bullets(page, dots: list[tuple[float, float]]) -> str:
    """Rebuild page text line by line, prepending "•" where a dot sits left."""
    lines = _group_lines(page.extract_words(x_tolerance=2))
    out: list[str] = []
    for ws in lines:
        ws.sort(key=lambda w: w["x0"])
        top = min(w["top"] for w in ws)
        x0 = ws[0]["x0"]
        txt = " ".join(w["text"] for w in ws)
        if any(abs(dt - top) <= 5 and dx < x0 and x0 - dx < 40 for dx, dt in dots):
            txt = "• " + txt
        out.append(txt)
    return "\n".join(out)


def _gutter(page) -> float | None:
    """X of a clear vertical column gutter, or None for a single column.

    Bins word centres across the page; if a near-empty vertical band sits in
    the middle with well-populated content on both sides, that band is the
    gutter between two columns.
    """
    return _detect_gutter(page.extract_words() or [], page.width)


def _detect_gutter(words: list[dict], w: float) -> float | None:
    """Pure column-gutter detection over word boxes (testable without a PDF)."""
    if len(words) < 40:
        return None
    nbins = 30
    binw = w / nbins
    counts = [0] * nbins
    for wd in words:
        c = (wd["x0"] + wd["x1"]) / 2
        counts[min(nbins - 1, int(c // binw))] += 1

    lo, hi = int(nbins * 0.35), int(nbins * 0.65)
    min_count, idx = min((counts[i], i) for i in range(lo, hi + 1))
    left = sum(counts[:idx])
    right = sum(counts[idx + 1:])
    total = len(words)
    peak = max(counts)
    gx = (idx + 0.5) * binw

    # Candidate gutter: near-empty band vs the densest column, content on both
    # sides. (Columns can be lopsided, so the side gate is low.)
    if not (min_count <= 0.15 * peak and left > total * 0.15 and right > total * 0.15):
        return None

    # Confirm a real column boundary, not a single column with a sparse middle
    # or a line-wrap. A true second column has its OWN lines: count lines living
    # entirely left vs entirely right of gx. Both sides must carry real lines.
    lines = _group_lines(words)
    left_only = sum(1 for ws in lines if all(w["x1"] < gx for w in ws))
    right_only = sum(1 for ws in lines if all(w["x0"] > gx for w in ws))
    need = max(3, int(len(lines) * 0.1))
    if left_only >= need and right_only >= need:
        return gx
    return None


def _group_lines(words: list[dict], tol: float = 3.0) -> list[list[dict]]:
    lines: list[list[dict]] = []
    for w in sorted(words, key=lambda w: w["top"]):
        if lines and abs(w["top"] - lines[-1][0]["top"]) <= tol:
            lines[-1].append(w)
        else:
            lines.append([w])
    return lines


def _extract_docx(path: Path) -> Document:
    import docx

    document = docx.Document(str(path))
    paras = [p.text for p in document.paragraphs if p.text.strip()]

    has_tables = len(document.tables) > 0
    # pull table cell text too so keyword/JD checks still see it
    for tbl in document.tables:
        for row in tbl.rows:
            for cell in row.cells:
                if cell.text.strip():
                    paras.append(cell.text)

    has_images = bool(document.inline_shapes)
    has_columns = _docx_has_columns(document)

    links = [
        rel.target_ref for rel in document.part.rels.values()
        if "hyperlink" in rel.reltype
    ]

    return Document(
        text=_normalize("\n".join(paras)), fmt="docx", source=path,
        has_tables=has_tables, has_columns=has_columns,
        has_images=has_images, is_scanned=False, links=_dedup(links),
    )


def _docx_has_columns(document) -> bool:
    # ponytail: read the section column count from sectPr; >1 means multi-column.
    try:
        for section in document.sections:
            cols = section._sectPr.find(
                "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}cols"
            )
            if cols is not None:
                num = cols.get(
                    "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}num"
                )
                if num and int(num) > 1:
                    return True
    except Exception:
        pass
    return False


def extract(path: str | Path) -> Document:
    """Extract text + layout flags from a PDF or DOCX resume."""
    path = Path(path)
    if not path.is_file():
        raise FileNotFoundError(f"No such file: {path}")
    suffix = path.suffix.lower()
    if suffix == ".pdf":
        return _extract_pdf(path)
    if suffix == ".docx":
        return _extract_docx(path)
    raise ValueError(f"Unsupported file type '{suffix}'. Use .pdf or .docx.")


def _selfcheck() -> None:
    # Round-trip a generated DOCX (no fixtures needed). PDF path is covered by
    # fixtures in the test suite (phase 8).
    import tempfile
    import docx

    with tempfile.TemporaryDirectory() as d:
        p = Path(d) / "r.docx"
        doc = docx.Document()
        doc.add_paragraph("Jane Doe")
        doc.add_paragraph("Experience: built things")
        doc.save(str(p))

        out = extract(p)
        assert out.fmt == "docx", out.fmt
        assert "Jane Doe" in out.text, out.text
        assert "built things" in out.text, out.text
        assert "\r" not in out.text
        assert out.is_scanned is False

    try:
        extract(Path(d) / "missing.docx")
    except FileNotFoundError:
        pass
    else:
        raise AssertionError("expected FileNotFoundError")

    try:
        extract("resume.txt")
    except (ValueError, FileNotFoundError):
        pass
    else:
        raise AssertionError("expected rejection of .txt")

    assert _normalize("Java​Script﻿") == "JavaScript", "zero-width strip"
    assert _normalize("a\x0d\x0ab") == "a\nb"
    assert _normalize("third-\nparty") == "third-party", "hyphen-break"

    # Column detection on synthetic word boxes (no PDF needed).
    def word(x0, x1, top):
        return {"x0": x0, "x1": x1, "top": top}

    # Offset, unequal-length columns (like a real layout, not aligned rows).
    two_col = ([word(30, 120, t) for t in range(0, 400, 12)]       # left column
               + [word(360, 450, t) for t in range(6, 250, 12)])   # right column
    assert _detect_gutter(two_col, 600) is not None, "missed two columns"

    one_col = [word(30, 500, t) for t in range(0, 600, 12)]       # full-width lines
    assert _detect_gutter(one_col, 600) is None, "false-split single column"

    print("extract selfcheck ok")


if __name__ == "__main__":
    _selfcheck()
