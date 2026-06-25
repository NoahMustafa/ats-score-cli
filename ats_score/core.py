"""Orchestration for V1: an ATS-readiness linter.

The overall score IS the ATS-readiness score — can a tracking system parse this
resume, and what is it missing. Writing advice (filler + AI tells) is reported
but does NOT affect the score. Content-quality grading (bullets/quantification)
is kept in the codebase but intentionally not wired into V1 — it was unreliable.

JD-match needs the embedding model, which is not bundled in V1. It runs only
when a JD is given AND the model is present on disk; otherwise the report just
lists the skills the parser could read (taxonomy-only, no model needed).
"""

from __future__ import annotations

from dataclasses import dataclass, field
from pathlib import Path

from .extract import extract, Document
from .checks_ats import check_ats, AtsResult
from .writing import check_writing, WritingResult
from .similarity import detect_skills
from .paths import bundled_path


def _model_available() -> bool:
    return bundled_path("data/potion-8M").exists()


@dataclass
class Report:
    overall: int                      # == ATS-readiness score in V1
    ats: AtsResult
    writing: WritingResult            # advice only, not scored
    detected_skills: list[str] = field(default_factory=list)
    similarity: object | None = None  # SimilarityResult when a JD + model exist
    jd_unavailable: bool = False      # JD given but model not bundled
    warnings: list[str] = field(default_factory=list)
    source: str = ""


def _read_jd(jd: str | Path) -> str:
    p = Path(jd)
    if p.is_file():
        return p.read_text(encoding="utf-8", errors="replace")
    return str(jd)  # allow passing the JD text directly


def score(resume_path: str | Path, jd_path: str | Path | None = None) -> Report:
    doc: Document = extract(resume_path)

    ats = check_ats(doc)
    writing = check_writing(doc)

    sim = None
    jd_unavailable = False
    detected: list[str] = []
    if jd_path is not None and _model_available():
        from .similarity import check_similarity
        sim = check_similarity(doc.text, _read_jd(jd_path))
    elif jd_path is not None:
        jd_unavailable = True          # model not in this build
        detected = detect_skills(doc.text)
    else:
        detected = detect_skills(doc.text)

    return Report(
        overall=ats.score, ats=ats, writing=writing,
        detected_skills=detected, similarity=sim,
        jd_unavailable=jd_unavailable, warnings=doc.warnings,
        source=str(resume_path),
    )


def _selfcheck() -> None:
    import tempfile
    import docx

    with tempfile.TemporaryDirectory() as d:
        p = Path(d) / "r.docx"
        doc = docx.Document()
        doc.add_paragraph("Jane Doe — jane@example.com — +1 555 123 4567 — Austin, TX")
        doc.add_paragraph("linkedin.com/in/janedoe")
        doc.add_paragraph("SUMMARY")
        doc.add_paragraph("Data engineer skilled in Python and SQL.")
        doc.add_paragraph("WORK EXPERIENCE")
        doc.add_paragraph("Built pipelines Jan 2024 to Mar 2024.")
        doc.add_paragraph("EDUCATION")
        doc.add_paragraph("BSc Computer Science")
        doc.add_paragraph("TECHNICAL SKILLS")
        doc.add_paragraph("Python, SQL, AWS, Airflow, ETL")
        doc.save(str(p))

        r = score(p)
        # Overall is exactly the ATS score in V1.
        assert r.overall == r.ats.score, (r.overall, r.ats.score)
        assert r.similarity is None
        assert r.detected_skills, "no JD should list detected skills"

        # JD path: either runs (model present) or reports unavailable — never
        # crashes. detected_skills stand in when the model is absent.
        rj = score(p, "python sql aws kubernetes")
        if _model_available():
            assert rj.similarity is not None and not rj.jd_unavailable
        else:
            assert rj.jd_unavailable and rj.similarity is None

    print("core selfcheck ok")


if __name__ == "__main__":
    _selfcheck()
