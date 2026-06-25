"""End-to-end V1: overall == ATS score, advice not scored, JD gated on model."""

import json
from pathlib import Path

import docx
import pytest

from ats_score.core import score, _model_available
from ats_score.report import to_dict, render_json, render


@pytest.fixture
def resume(tmp_path) -> Path:
    p = tmp_path / "r.docx"
    d = docx.Document()
    d.add_paragraph("Jane Doe — jane@example.com — +1 555 123 4567 — Austin, TX")
    d.add_paragraph("linkedin.com/in/janedoe")
    d.add_paragraph("SUMMARY")
    d.add_paragraph("Data engineer skilled in Python and SQL.")
    d.add_paragraph("WORK EXPERIENCE")
    d.add_paragraph("Built pipelines Jan 2024 to Mar 2024.")
    d.add_paragraph("EDUCATION")
    d.add_paragraph("BSc Computer Science")
    d.add_paragraph("TECHNICAL SKILLS")
    d.add_paragraph("Python, SQL, AWS, Airflow, ETL")
    d.save(str(p))
    return p


def test_overall_equals_ats_score(resume):
    r = score(resume)
    assert r.overall == r.ats.score
    assert 0 <= r.overall <= 100


def test_no_jd_lists_skills(resume):
    r = score(resume)
    assert r.similarity is None
    assert r.detected_skills


def test_jd_gated_on_model(resume):
    r = score(resume, "python sql aws kubernetes")
    if _model_available():
        assert r.similarity is not None and not r.jd_unavailable
    else:
        assert r.jd_unavailable and r.similarity is None


def test_writing_advice_does_not_affect_score(resume):
    # Overall is purely the ATS score; advice findings never change it.
    r = score(resume)
    assert r.overall == r.ats.score


def test_json_shape(resume):
    d = to_dict(score(resume))
    assert d["overall"] == d["ats"]["score"]
    assert "writing_advice" in d
    assert "content" not in d            # content cut in V1
    assert json.loads(render_json(score(resume)))["overall"] == d["overall"]


def test_render_plain_when_not_a_tty(resume):
    out = render(score(resume))
    assert "\x1b[" not in out
    assert "ATS readiness:" in out


def test_unsupported_file_type_rejected(tmp_path):
    p = tmp_path / "resume.txt"
    p.write_text("hi", encoding="utf-8")
    with pytest.raises(ValueError):
        score(p)
